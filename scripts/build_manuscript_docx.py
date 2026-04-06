#!/usr/bin/env python3
"""Build a single formatted Word document from manuscript components."""

import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import re

ROOT = Path(__file__).parent.parent
MANUSCRIPT = ROOT / "manuscript"
OUTPUTS = ROOT / "outputs"

doc = Document()

# ── Page Setup ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Styles ──────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0  # Double spacing
style.paragraph_format.space_after = Pt(0)

# Heading styles
for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Times New Roman'
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.font.bold = True
    h.paragraph_format.line_spacing = 2.0
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)

doc.styles['Heading 1'].font.size = Pt(14)
doc.styles['Heading 2'].font.size = Pt(12)
doc.styles['Heading 3'].font.size = Pt(11)


def add_paragraph(text, bold=False, italic=False, style_name='Normal', alignment=None):
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    return p


def add_table(headers, rows, caption=None, footnote=None):
    """Add a formatted table with caption and footnote."""
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(10)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'

    if footnote:
        p = doc.add_paragraph()
        run = p.add_run(footnote)
        run.italic = True
        run.font.size = Pt(9)

    doc.add_paragraph()  # spacing
    return table


def add_figure(image_path, caption, width=Inches(5.5)):
    """Add a figure with caption."""
    if Path(image_path).exists():
        doc.add_picture(str(image_path), width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph()
        run = p.add_run(f"[FIGURE PLACEHOLDER: {image_path}]")
        run.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    run = p.add_run(caption)
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "Optimising the timing and targeting of interventions to reduce DTP "
    "vaccination dropout using offline reinforcement learning: a four-stage "
    "analytical framework applied to the Nigeria DHS 2024"
)
run.bold = True
run.font.size = Pt(16)

doc.add_paragraph()

add_paragraph(
    "Uthman OA et al.",
    bold=True,
    alignment=WD_ALIGN_PARAGRAPH.CENTER
)
add_paragraph(
    "Warwick Applied Health, Warwick Medical School, University of Warwick",
    italic=True,
    alignment=WD_ALIGN_PARAGRAPH.CENTER
)
doc.add_paragraph()
add_paragraph(
    "Corresponding author: Professor Olalekan A. Uthman",
    alignment=WD_ALIGN_PARAGRAPH.CENTER
)
add_paragraph(
    "[Affiliations and email to be completed]",
    italic=True,
    alignment=WD_ALIGN_PARAGRAPH.CENTER
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════
doc.add_heading('Abstract', level=1)

p = doc.add_paragraph()
run = p.add_run('Background: ')
run.bold = True
p.add_run(
    "Nigeria accounts for the largest absolute burden of under-vaccinated children globally, "
    "with persistent dropout between DTP1 and DTP3 remaining a critical barrier to full immunisation "
    "coverage despite improving trends. Conventional dropout interventions — SMS reminders, community "
    "health worker visits, defaulter tracing — are deployed uniformly without considering optimal "
    "sequencing or personalisation across the vaccination schedule. No study has applied reinforcement "
    "learning to vaccination dropout. We applied a four-stage analytical framework to identify which "
    "children to target, with which intervention, and at which dose transition, using nationally "
    "representative survey data."
)

p = doc.add_paragraph()
run = p.add_run('Methods and Findings: ')
run.bold = True
p.add_run(
    "We analysed 3,194 children aged 12–23 months who had received DTP1 in the Nigeria Demographic "
    "and Health Survey 2024. The overall DTP1-to-DTP3 dropout rate was 14.8%, with geographic "
    "heterogeneity ranging from 8.6% (South South) to 19.6% (North Central). Transition-specific "
    "XGBoost models, trained with 200-trial Bayesian hyperparameter optimisation and cluster-robust "
    "cross-validation, achieved AUC-ROC of 0.910 (95% CI: 0.885–0.932) for DTP1-to-DTP2 dropout "
    "and 0.943 (0.929–0.957) for DTP2-to-DTP3 dropout, significantly outperforming logistic "
    "regression (DeLong p < 1 × 10⁻⁹). Contrary to the hypothesis that static sociodemographic "
    "factors would dominate, SHAP decomposition by Andersen's Behavioural Model revealed that "
    "dynamic temporal features — inter-dose intervals, delay accumulation, community dropout "
    "rates — accounted for 35.9% and 59.0% of total feature importance at transitions 1 and 2, "
    "respectively. A conservative Q-learning offline reinforcement learning policy (α = 5.0) "
    "achieved 6.1–6.7% improvement over the observed behaviour policy. In microsimulation "
    "(10,000 children, 1,000 bootstrap iterations), risk-targeted allocation was the most "
    "cost-effective non-dominant strategy (DTP3: 88.2%, ICER: ₦8,007 per additional completion), "
    "capturing 92% of the gain from universal community health worker deployment at 35% of the "
    "cost. The reinforcement learning-optimised policy achieved equivalent DTP3 completion (88.2%) "
    "at 2.6 times the cost (₦903 vs ₦341 per child). All scenarios reduced the poorest-to-richest "
    "equity gap from 7.8 to 6.6–7.3 percentage points."
)

p = doc.add_paragraph()
run = p.add_run('Conclusions: ')
run.bold = True
p.add_run(
    "Dynamic vaccination trajectory features — not static demographics — are the strongest "
    "predictors of dropout at both dose transitions, supporting trajectory-responsive intervention "
    "design. Risk-targeted allocation guided by machine-learning risk scores is the most "
    "cost-effective strategy for Nigeria. Offline reinforcement learning provides a novel framework "
    "for sequential intervention optimisation in immunisation, though it did not outperform static "
    "targeting over a two-transition horizon. Applications to longer vaccine schedules with "
    "prospective intervention data may better demonstrate the sequential advantage of reinforcement "
    "learning."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# INTRODUCTION
# ══════════════════════════════════════════════════════════════
doc.add_heading('Introduction', level=1)

intro_paras = [
    "Vaccination dropout, defined as the failure to complete a multi-dose vaccine series after "
    "initiating it, represents a distinct programmatic challenge from both zero-dose children who "
    "never access immunisation services and missed opportunities for vaccination where children "
    "attend health facilities but are not vaccinated [1,2]. Children who receive DTP1 but do not "
    "complete the three-dose series have successfully entered the immunisation cascade, indicating "
    "that the health system engaged them at least once, yet failed to retain them through the "
    "schedule [3]. Globally, the DTP1-to-DTP3 dropout rate serves as a key performance indicator "
    "for immunisation programmes, with the World Health Organization defining it as the proportional "
    "difference between first and third dose coverage [4]. In sub-Saharan Africa, the weighted "
    "pentavalent dropout rate is approximately 20.9%, with Western Africa bearing the highest "
    "burden at 46.0% [5].",

    "In Nigeria, DTP series dropout has declined from 46% in 2003 to approximately 28% in 2018, "
    "yet remains a substantial barrier to achieving full immunisation coverage [6]. National DTP3 "
    "coverage was estimated at 53.5% in the 2024 Nigeria Demographic and Health Survey, with marked "
    "geographical disparities: the North-West and North-East geopolitical zones consistently exhibit "
    "the lowest coverage and highest dropout rates [7]. Prior analyses of dropout determinants using "
    "Demographic and Health Survey data have identified maternal education, household wealth, "
    "antenatal care attendance, place of delivery, and urban-rural residence as consistent "
    "predictors, typically framed within Andersen's Behavioural Model of Health Services Use "
    "[1,5,8]. However, these analyses treat dropout as a single binary outcome, ignoring the "
    "sequential nature of the vaccination schedule and the possibility that distinct predictor "
    "profiles operate at different cascade transitions.",

    "Current interventions to reduce dropout, including SMS reminders, community health worker home "
    "visits, facility-based defaulter tracing, and conditional incentives, have demonstrated "
    "effectiveness in randomised trials, with effect sizes ranging from 5–15% for SMS reminders to "
    "15–25% for community health worker visits [9,10,11]. However, these interventions are typically "
    "deployed uniformly across the vaccination schedule without considering the optimal sequencing, "
    "timing, or targeting of different intervention types at different dose steps for different "
    "children. A caregiver whose child missed DTP2 at 10 weeks may require a different intervention "
    "from one whose child received DTP1 and DTP2 on time but fails to return for DTP3. This "
    "sequential decision problem, where the optimal action depends on the child's evolving state "
    "within the immunisation cascade, is naturally suited to reinforcement learning [12,13]. Offline "
    "reinforcement learning methods, which learn optimal policies from fixed observational datasets "
    "without requiring real-time experimentation, have been successfully applied to clinical "
    "decision-making in sepsis management and mechanical ventilation [14,15,16], but have not been "
    "applied to vaccination dropout.",

    "This study aimed to (1) construct transition-specific prediction models for DTP dropout using "
    "XGBoost with Andersen-stratified SHAP decomposition to identify distinct predictor profiles at "
    "each cascade transition; (2) formulate the immunisation schedule as a Markov decision process "
    "and learn optimal sequential intervention policies using offline reinforcement learning; "
    "(3) develop a contextual multi-armed bandit for budget-constrained community-level intervention "
    "allocation; and (4) validate the reinforcement learning-derived policy through microsimulation "
    "comparing it against standard static intervention approaches."
]

for para in intro_paras:
    doc.add_paragraph(para)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# METHODS (from 01_methods.md — key sections)
# ══════════════════════════════════════════════════════════════
doc.add_heading('Methods', level=1)

# Read methods markdown and convert key paragraphs
methods_text = (MANUSCRIPT / "01_methods.md").read_text()

# Extract paragraphs between headers
import re

def extract_section(text, start_header, end_header=None):
    """Extract text between two markdown headers."""
    pattern = rf'^## {re.escape(start_header)}$(.*?)'
    if end_header:
        pattern += rf'(?=^## {re.escape(end_header)}$)'
    else:
        pattern += r'(?=^## |\Z)'
    match = re.search(pattern, text, re.MULTILINE | re.DOTALL)
    if match:
        return match.group(1).strip()
    return ""

# Write methods sections
methods_sections = [
    ("Study design and data source", 2),
    ("Study population", 2),
    ("Outcome definitions", 2),
    ("Predictor variables", 2),
    ("Analytical framework", 2),
    ("Survey design and statistical analysis", 2),
    ("Software", 2),
]

for section_title, level in methods_sections:
    content = extract_section(methods_text, section_title)
    if not content:
        continue

    doc.add_heading(section_title, level=level)

    # Split into sub-sections and paragraphs
    lines = content.split('\n')
    current_para = []

    for line in lines:
        line = line.strip()
        if not line:
            if current_para:
                doc.add_paragraph(' '.join(current_para))
                current_para = []
        elif line.startswith('### '):
            if current_para:
                doc.add_paragraph(' '.join(current_para))
                current_para = []
            doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('#### '):
            if current_para:
                doc.add_paragraph(' '.join(current_para))
                current_para = []
            p = doc.add_paragraph()
            run = p.add_run(line.replace('#### ', ''))
            run.bold = True
        elif line.startswith('- **') or line.startswith('- S'):
            if current_para:
                doc.add_paragraph(' '.join(current_para))
                current_para = []
            # Convert to paragraph with bold label
            doc.add_paragraph(line.lstrip('- '))
        elif line.startswith('#') or line.startswith('```') or line.startswith('|'):
            continue  # Skip code blocks and tables in methods
        else:
            current_para.append(line)

    if current_para:
        doc.add_paragraph(' '.join(current_para))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# RESULTS
# ══════════════════════════════════════════════════════════════
doc.add_heading('Results', level=1)

# Load the merged results
results_text = (MANUSCRIPT / "03_results_merged.md").read_text()

results_sections = [
    "Study population and immunisation cascade",
    "Descriptive characteristics",
    "Spatial variation in dropout",
    "Prediction model discrimination",
    "Prediction model calibration",
    "SHAP feature importance and Andersen domain decomposition",
    "Offline reinforcement learning",
    "Community-level bandit allocation",
    "Microsimulation: DTP3 completion, cost-effectiveness, and equity",
    "RL-optimised versus static targeting",
]

for section_title in results_sections:
    content = extract_section(results_text, section_title)
    if not content:
        continue

    doc.add_heading(section_title, level=2)

    lines = content.split('\n')
    current_para = []
    for line in lines:
        line = line.strip()
        if not line:
            if current_para:
                doc.add_paragraph(' '.join(current_para))
                current_para = []
        elif line.startswith('#'):
            continue
        elif line.startswith('|'):
            continue  # Tables added separately
        else:
            current_para.append(line)
    if current_para:
        doc.add_paragraph(' '.join(current_para))

# ── Insert Tables ──────────────────────────────────────────

# Table 2: Model Performance
doc.add_paragraph()
add_table(
    headers=["Metric", "Model T1", "Model T2"],
    rows=[
        ["N", "3,194", "3,023"],
        ["Event rate", "4.0%", "3.9%"],
        ["AUC-ROC (95% CI)", "0.910 (0.885–0.932)", "0.943 (0.929–0.957)"],
        ["AUC-PR (95% CI)", "0.299 (0.230–0.387)", "0.548 (0.454–0.640)"],
        ["Brier score (recalibrated)", "0.032 (0.036)", "0.026 (0.027)"],
        ["Calibration slope (recalibrated)", "0.959 (0.874)", "0.969 (1.596)"],
        ["Calibration intercept (recalibrated)", "−0.065 (−0.661)", "−0.104 (0.724)"],
        ["DeLong p (vs LR)", "8.1 × 10⁻¹⁰", "9.7 × 10⁻¹⁰"],
        ["LR baseline AUC-ROC", "0.818", "0.863"],
    ],
    caption="Table 2. XGBoost model performance across cascade transitions.",
    footnote="AUC-ROC = area under the receiver operating characteristic curve. AUC-PR = area under the "
             "precision-recall curve. Values in parentheses are pre-recalibration metrics. Isotonic regression "
             "used for post-hoc recalibration. Cross-validation used cluster-robust folds holding out entire "
             "primary sampling units. LR = logistic regression."
)

# Table 3: Andersen SHAP
add_table(
    headers=["Domain (n features)", "T1 mean |SHAP| (%)", "T2 mean |SHAP| (%)", "Full mean |SHAP| (%)"],
    rows=[
        ["Predisposing (16)", "0.79 (17.3%)", "0.21 (10.1%)", "0.90 (17.1%)"],
        ["Enabling (20)", "0.90 (19.7%)", "0.23 (11.2%)", "0.80 (15.2%)"],
        ["Need (11)", "1.24 (27.1%)", "0.41 (19.7%)", "0.82 (15.6%)"],
        ["Dynamic (8)", "1.63 (35.9%)", "1.24 (59.0%)", "2.74 (52.1%)"],
    ],
    caption="Table 3. Andersen domain SHAP decomposition across cascade transitions.",
    footnote="SHAP = SHapley Additive exPlanations. Values computed on held-out test set using TreeExplainer. "
             "Percentages represent domain share of total mean absolute SHAP importance. Dynamic features "
             "include inter-dose intervals, delay accumulation, community dropout rates, child age, and dose count."
)

# Table 4: OPE
add_table(
    headers=["Metric", "Value"],
    rows=[
        ["Behaviour policy value", "1.104"],
        ["WIS policy value", "1.171"],
        ["FQE policy value", "1.178"],
        ["WIS improvement", "+6.1%"],
        ["FQE improvement", "+6.7%"],
        ["OOD action frequency (CQL)", "33.3%"],
        ["CQL α", "5.0"],
        ["Validation episodes", "1,569"],
    ],
    caption="Table 4. Off-policy evaluation of the CQL-learned intervention policy.",
    footnote="WIS = weighted importance sampling (ε = 0.1). FQE = fitted Q-evaluation. "
             "OOD = out-of-distribution (actions observed in <1% of training data). "
             "CQL = conservative Q-learning."
)

# Table 5: Microsimulation
add_table(
    headers=["Scenario", "DTP3 rate (95% CI)", "Cost/child (₦)", "ICER vs S0 (₦)", "Equity gap (pp)"],
    rows=[
        ["S0: Status quo", "85.9 (85.2–86.6)", "155", "—", "7.8"],
        ["S1: Uniform SMS", "87.1 (86.4–87.7)", "98", "Dominant", "7.2"],
        ["S2: Uniform CHW", "88.4 (87.8–89.0)", "979", "32,742", "6.6"],
        ["S3: Risk-targeted", "88.2 (87.6–88.8)", "341", "8,007", "6.7"],
        ["S4: RL-optimised", "88.2 (87.5–88.8)", "903", "32,434", "6.8"],
        ["S5: Bandit-allocated", "87.1 (86.5–87.8)", "98", "Dominant", "7.3"],
    ],
    caption="Table 5. Microsimulation scenario comparison (10,000 children, 1,000 bootstrap iterations).",
    footnote="ICER = incremental cost-effectiveness ratio relative to S0. Dominant = better outcomes at "
             "lower cost. Equity gap = absolute difference in DTP3 rate between poorest and richest wealth "
             "quintiles (percentage points). Effect sizes applied as relative risk reductions: "
             "a₁=10%, a₂=20%, a₃=25%, a₄=14%."
)

doc.add_page_break()

# ── Insert Figures ──────────────────────────────────────────

doc.add_heading('Figures', level=1)

# Figure 1 — placeholder
add_figure(
    "PLACEHOLDER",
    "Figure 1. Four-stage analytical framework for optimising sequential interventions to reduce "
    "DTP vaccination dropout. Stage 1: immunisation cascade characterisation and transition-specific "
    "XGBoost with SHAP/Andersen decomposition. Stage 2: offline reinforcement learning (FQI + CQL). "
    "Stage 3: contextual multi-armed bandit allocation. Stage 4: microsimulation validation."
)

# Figure 2
add_figure(
    OUTPUTS / "stage1" / "cascade_by_zone.png",
    "Figure 2. DTP immunisation cascade by geopolitical zone, Nigeria DHS 2024. Bar heights represent "
    "the proportion of DTP1 recipients who received each subsequent dose. WHO dropout rate annotated "
    "for each zone. South South had the lowest dropout (8.6%); North Central the highest (19.6%)."
)

# Figure 3
add_figure(
    OUTPUTS / "stage0" / "dropout_choropleth_map.png",
    "Figure 3A. Choropleth map of DTP1-to-DTP3 dropout prevalence by state. Darker shading indicates "
    "higher dropout rates. Global Moran's I = −0.044 (p = 0.452), indicating no significant spatial "
    "autocorrelation."
)

add_figure(
    OUTPUTS / "stage0" / "local_moran_clusters_map.png",
    "Figure 3B. Local indicators of spatial association (LISA) cluster map. Only 4 of 37 states showed "
    "significant local clustering (1 High-High, 3 High-Low outliers; 33 not significant)."
)

# Figure 4
add_figure(
    OUTPUTS / "stage0" / "dropout_funnel_plot.png",
    "Figure 4. Funnel plot of state-level DTP dropout prevalence against sample size. Dashed lines: "
    "95% control limits; dotted lines: 99.7% control limits. States outside limits exhibit "
    "statistically unusual dropout rates."
)

# Figure 5
add_figure(
    OUTPUTS / "stage1" / "andersen_decomp_comparison.png",
    "Figure 5. Andersen domain SHAP decomposition comparison across cascade transitions. Dynamic "
    "temporal features dominated all models: T1 = 35.9%, T2 = 59.0%, Full = 52.1%. The a priori "
    "hypothesis of predisposing dominance at T1 and enabling/need dominance at T2 was not supported."
)

# Figure 6
for model in ["t1", "t2"]:
    label = "T1 (DTP1→DTP2)" if model == "t1" else "T2 (DTP2→DTP3)"
    add_figure(
        OUTPUTS / "stage1" / f"roc_pr_model_{model}.png",
        f"Figure 6{'A' if model=='t1' else 'B'}. ROC and precision-recall curves for Model {label}."
    )
    add_figure(
        OUTPUTS / "stage1" / f"calibration_model_{model}_recalibrated.png",
        f"Figure 6{'C' if model=='t1' else 'D'}. Calibration curves for Model {label} before and "
        "after isotonic regression recalibration."
    )

# Figure 7
add_figure(
    OUTPUTS / "stage2" / "cql_analysis.png",
    "Figure 7A. Action distribution comparison between inferred behaviour policy and CQL-learned "
    "policy (α = 5.0). The CQL policy shifts from SMS-dominated (70.6%) to a diversified mix: "
    "CHW 29.4%, incentive 22.9%, no intervention 20.0%, SMS 17.3%, recall 10.4%."
)

add_figure(
    OUTPUTS / "stage2" / "q_values_by_action.png",
    "Figure 7B. Distribution of Q-values by action, showing estimated long-term value of each "
    "intervention type."
)

# Figure 8
add_figure(
    OUTPUTS / "stage3" / "microsim_scenarios.png",
    "Figure 8. Microsimulation scenario comparison of DTP3 completion rates with 95% confidence "
    "intervals (1,000 bootstrap iterations, N = 10,000). All scenarios improved on the status quo "
    "(85.9%). S3 (risk-targeted) was the most cost-effective non-dominant strategy (ICER ₦8,007)."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# DISCUSSION
# ══════════════════════════════════════════════════════════════
doc.add_heading('Discussion', level=1)

discussion_text = (MANUSCRIPT / "04_discussion.md").read_text()

disc_sections = [
    "Principal findings",
    "Comparison with previous studies",
    "Policy implications",
    "Strengths and limitations",
    "Conclusion",
]

for section_title in disc_sections:
    content = extract_section(discussion_text, section_title)
    if not content:
        continue

    doc.add_heading(section_title, level=2)

    # Handle sub-subsections
    lines = content.split('\n')
    current_para = []
    for line in lines:
        line = line.strip()
        if not line:
            if current_para:
                doc.add_paragraph(' '.join(current_para))
                current_para = []
        elif line.startswith('### '):
            if current_para:
                doc.add_paragraph(' '.join(current_para))
                current_para = []
            doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('#') or line.startswith('---'):
            continue
        else:
            current_para.append(line)
    if current_para:
        doc.add_paragraph(' '.join(current_para))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════
doc.add_heading('References', level=1)

# Write a simplified numbered reference list
refs = [
    "1. Uthman OA, Sambala EZ, et al. Systematic review of social determinants of childhood immunisation in LMICs and equity impact analysis in Nigeria. PLOS ONE. 2024;19(2):e0297326.",
    "2. Sridhar S, Maleq N, Guillermet E, et al. A systematic literature review of missed opportunities for immunization in LMICs. Vaccine. 2014;32(51):6870-6879.",
    "3. Adedokun ST, et al. Zero-dose, under-immunized, and dropout children in Nigeria: trend and contributing factors. Vaccines. 2023;11(1):167.",
    "4. WHO/UNICEF. Immunisation Agenda 2030: A Global Strategy to Leave No One Behind. Geneva: WHO; 2020.",
    "5. Tessema ZT, et al. Individual and community-level determinants of pentavalent vaccination dropouts in SSA: multilevel analysis. BMC Public Health. 2024;24:456.",
    "6. Adedokun ST, et al. Zero-dose, under-immunized, and dropout children in Nigeria. Vaccines. 2023;11(1):167.",
    "7. Adeyinka DA, et al. Widening geographical inequities in DTP vaccination coverage across Nigeria (2018-2024). Vaccines. 2025;13(11):1135.",
    "8. Andersen RM. Revisiting the behavioral model and access to medical care. J Health Soc Behav. 1995;36(1):1-10.",
    "9. Eze P, et al. Effect of mHealth interventions on childhood immunization in LMICs: meta-analysis. Hum Vaccin Immunother. 2021;17(2):507-515.",
    "10. Brugha RF, Kevany JP. Maximizing immunization coverage through home visits: controlled trial in Ghana. Bull WHO. 1996;74(5):517-524.",
    "11. Banerjee AV, Duflo E, et al. Improving immunisation coverage in rural India: clustered RCT with and without incentives. BMJ. 2010;340:c2220.",
    "12. Ernst D, Geurts P, Wehenkel L. Tree-based batch mode reinforcement learning. JMLR. 2005;6:503-556.",
    "13. Kumar A, Zhou A, Tucker G, Levine S. Conservative Q-learning for offline RL. NeurIPS. 2020;33:1179-1191.",
    "14. Komorowski M, Celi LA, et al. The AI Clinician learns optimal treatment strategies for sepsis. Nat Med. 2018;24(11):1716-1720.",
    "15. Raghu A, et al. Deep reinforcement learning for sepsis treatment. arXiv:1711.09602. 2017.",
    "16. Raghu A, Komorowski M, Singh S. Model-based RL for sepsis treatment. arXiv:1811.09602. 2018.",
    "17. Collins GS, et al. TRIPOD+AI statement: updated guidance for reporting prediction models. BMJ. 2024;385:e078378.",
    "18. NPC Nigeria and ICF. Nigeria Demographic and Health Survey 2024. Abuja/Rockville: NPC/ICF; 2025.",
    "19. ICF. The DHS Program Geospatial Covariates Dataset. Rockville: ICF; 2023.",
    "20. Andersen RM. Revisiting the behavioral model and access to medical care. J Health Soc Behav. 1995;36(1):1-10.",
    "21. NPHCDA. National Strategic Immunisation Plan and Schedule of Services 2018-2028. Abuja: NPHCDA; 2018.",
    "22. Chen T, Guestrin C. XGBoost: a scalable tree boosting system. KDD. 2016:785-794.",
    "23. Akiba T, et al. Optuna: next-generation hyperparameter optimization framework. KDD. 2019:2623-2631.",
    "24. Efron B, Tibshirani RJ. An Introduction to the Bootstrap. Chapman and Hall/CRC; 1993.",
    "25. DeLong ER, et al. Comparing areas under correlated ROC curves: nonparametric approach. Biometrics. 1988;44(3):837-845.",
    "26. Spiegelhalter DJ. Probabilistic prediction in patient management. Stat Med. 1986;5(5):421-433.",
    "27. Niculescu-Mizil A, Caruana R. Predicting good probabilities with supervised learning. ICML. 2005:625-632.",
    "28. Lundberg SM, Lee SI. A unified approach to interpreting model predictions. NeurIPS. 2017;30:4765-4774.",
    "29. Banerjee AV, et al. Improving immunisation coverage in rural India: clustered RCT. BMJ. 2010;340:c2220.",
    "30. Ernst D, et al. Tree-based batch mode RL. JMLR. 2005;6:503-556.",
    "31. Kumar A, et al. Conservative Q-learning for offline RL. NeurIPS. 2020;33:1179-1191.",
    "32. Raghu A, et al. Model selection for offline RL: practical considerations for healthcare. MLHC. 2022.",
    "33. Li L, et al. Contextual-bandit approach to personalized recommendation. WWW. 2010:661-670.",
    "34. Anselin L. Local indicators of spatial association — LISA. Geogr Anal. 1995;27(2):93-115.",
    "35. Komorowski M, et al. The AI Clinician for sepsis. Nat Med. 2018;24(11):1716-1720.",
    "36. Wang Z, et al. K-CQL: CQL-based ventilator management. arXiv. 2025.",
    "37. Nambiar M, et al. PROP-RL: practical pipeline for offline RL in clinical workflow. arXiv. 2025.",
    "38. Marseille E, et al. Thresholds for cost-effectiveness of interventions. Bull WHO. 2015;93(2):118-124.",
    "39. Chakraborty B, Moodie EEM. Statistical Methods for Dynamic Treatment Regimes. Springer; 2013.",
    "40. Tsiatis AA, et al. Dynamic Treatment Regimes: Statistical Methods for Precision Medicine. CRC Press; 2020.",
    "41. WHO. Recommendations on Digital Interventions for Health System Strengthening. Geneva: WHO; 2019.",
    "42. Adegboye OA, et al. Factors influencing childhood immunisation uptake in Africa: systematic review. BMC Public Health. 2021;21:1475.",
    "43. Wiysonge CS, et al. Individual and contextual factors associated with low childhood immunisation in SSA. PLOS ONE. 2012;7(5):e37905.",
    "44. Faye CM, et al. Immunization coverage and determinants in East Africa: Bayesian hierarchical model. BMC Public Health. 2024.",
    "45. Laryea DO, et al. Determinants of uptake of third doses of OPV and DTP in Ibadan Nigeria. Int Health. 2014;6(3):213-224.",
    "46. Mekonnen ZA, et al. Mobile phone-based interventions for childhood immunization in SSA: systematic review. BMC Public Health. 2021;21:2200.",
    "47. Sato Y, Takasaki Y. Offline reinforcement learning for public health: a tutorial. Stat Med. 2024;43(5):1012-1030.",
    "48. Zhang Y, et al. Methods in DTRs using observational healthcare data: systematic review. Comput Methods Programs Biomed. 2025.",
    "49. Jacobson Vann JC, et al. Patient reminder and recall interventions to improve immunization rates. Cochrane Database Syst Rev. 2018;1:CD003941.",
    "50. Gavi. Zero-Dose Learning Agenda. Geneva: Gavi; 2022.",
    "51. ICF. The DHS Program: DHS Methodology. Rockville: ICF; 2023.",
    "52. Levine E, Tsiatis AA, Schulte PJ. Dynamic treatment regimes. In: Handbook of Missing Data Methodology. CRC; 2014.",
    "53. Yechezkel M, et al. Multi-city dynamic allocation of COVID-19 vaccines using RL. Nat Commun. 2023;14:3960.",
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.size = Pt(10)

# ── Save ──────────────────────────────────────────────────
output_path = MANUSCRIPT / "manuscript_full.docx"
doc.save(str(output_path))
print(f"Saved to: {output_path}")
print(f"Size: {output_path.stat().st_size / 1024:.0f} KB")
